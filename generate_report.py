from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

def set_run_style(run, font_name='Times New Roman', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_image_placeholder(doc, text, height_inches=3.5):
    """Adds a bordered placeholder for images to increase page count and guide the user."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(f"\n[ INSERT IMAGE HERE ]\n{text}\n")
    runner.bold = True
    runner.font.color.rgb = RGBColor(128, 128, 128)
    runner.font.size = Pt(11)
    # Add some spacing
    doc.add_paragraph()

def main():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Government Residence Women Polytechnic, Tasgaon\n")
    set_run_style(run, font_size=16, bold=True)
    
    run = p.add_run("2026 – 2027\n\n")
    set_run_style(run, font_size=14, bold=True)
    
    run = p.add_run("A Report on\n“INTERNSHIP”\n(ITR – 315004)\n\n")
    set_run_style(run, font_size=16, bold=True)
    
    run = p.add_run("Project Title:\nAI College Assistant + Study Planner\n\n")
    set_run_style(run, font_size=14, bold=True, color=RGBColor(0, 51, 102))
    
    run = p.add_run("Department of Computer Engineering\n")
    set_run_style(run, font_size=14, bold=True)
    
    run = p.add_run("MAHARASHTRA STATE BOARD OF TECHNICAL EDUCATION, MUMBAI\n\n\n\n")
    set_run_style(run, font_size=14, bold=True)
    
    run = p.add_run("Submitted by:\nAmruta Chandrakant Shinde\nRoll No: 30 | Enrolment No: 24210280163\n\n")
    set_run_style(run, font_size=12, bold=True)
    
    run = p.add_run("Under the Guidance of:\nDr. Vinod R. Falmari (Mentor)\nSmt. S. B. Patil (HOD, Computer Engineering)\nDr. S. B. Patil (Principal, GRWPT)")
    set_run_style(run, font_size=12, bold=True)

    doc.add_page_break()

    # =========================================================================
    # CERTIFICATE
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CERTIFICATE")
    set_run_style(run, font_size=16, bold=True)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("This is to certify that the student mentioned below of Diploma in Computer Engineering of Government Residence Women Polytechnic, Tasgaon (1228), has satisfactorily completed a 12-week duration internship and has submitted this report as partial fulfilment of the prescribed curriculum of the Maharashtra State Board of Technical Education, Mumbai, for the academic year 2026–2027.")
    set_run_style(run, font_size=12)

    # Certificate Table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Roll No."
    hdr[1].text = "Enrolment No."
    hdr[2].text = "Name of Student"
    
    row1 = table.rows[1].cells
    row1[0].text = "30"
    row1[1].text = "24210280163"
    row1[2].text = "Amruta Chandrakant Shinde"

    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    table_sig = doc.add_table(rows=1, cols=3)
    table_sig.cell(0, 0).text = "Dr. Vinod R. Falmari\nMentor"
    table_sig.cell(0, 1).text = "Smt. S. B. Patil\nHead of Department, Computer Engineering"
    table_sig.cell(0, 2).text = "Dr. S. B. Patil\nPrincipal, GRWPT"
    
    for cell in table_sig.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =========================================================================
    # ABSTRACT
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ABSTRACT")
    set_run_style(run, font_size=16, bold=True)
    doc.add_paragraph()

    abstract_text = (
        "The 12-week industrial training at iGAP Technologies Pvt. Ltd. provided an excellent opportunity to bridge the gap between academic learning and industrial practices. "
        "The training was designed to enhance technical skills, professional competence, and practical exposure to real-world IT projects. During the program, I gained hands-on experience "
        "in areas such as software development, data science, artificial intelligence, and machine learning using Python. The training emphasized key concepts including data preprocessing, "
        "exploratory data analysis, implementation of machine learning models, and deployment of solutions for real-time applications.\n\n"
        "Industrial Training provides required professional and practical skills to the students. It is an essential part in the development of the practical and professional skills required "
        "of an engineer and an aid to prospective employment. For this training, I joined iGAP Technologies Pvt. Ltd., where I learned Data Science and AI/ML using Python and developed "
        "a comprehensive project.\n\n"
        "The project developed is the 'AI College Assistant + Study Planner'. Students often struggle with lengthy educational materials like college notes, textbooks, and PDF documents. "
        "Finding specific topics or answers manually is time-consuming. Traditional PDF readers only allow keyword searches and lack contextual understanding, automatic summarization, "
        "or practice question generation. The AI College Assistant solves this by providing an intelligent platform where students can upload college PDFs and interact with the material "
        "using natural language. The system extracts text, divides it into semantic chunks, generates embeddings using HuggingFace models, and stores them in a FAISS vector database. "
        "Using Retrieval-Augmented Generation (RAG) powered by the Google Gemini AI model, the system provides accurate, context-aware answers, generates concise study summaries, "
        "creates practice MCQs, identifies important examination questions, and offers a personalized AI-driven study planner with progress tracking. This project successfully demonstrates "
        "the practical application of Generative AI and NLP in solving real-world educational challenges."
    )
    p = doc.add_paragraph(abstract_text)
    set_run_style(p.runs[0], font_size=12)
    doc.add_page_break()

    # =========================================================================
    # ACKNOWLEDGEMENT
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ACKNOWLEDGEMENT")
    set_run_style(run, font_size=16, bold=True)
    doc.add_paragraph()

    ack_text = (
        "I would like to express my sincere gratitude to iGAP Technologies Pvt. Ltd. for providing me with the opportunity to undergo 12 weeks of industrial training in Data Science "
        "and Artificial Intelligence/Machine Learning using Python. This training has been an invaluable learning experience, allowing me to apply theoretical knowledge to practical, "
        "real-world problems and gain industry-relevant skills.\n\n"
        "It is a great privilege for me to express my sincere thanks to Dr. S. B. Patil, Principal of Government Residence Women Polytechnic, Tasgaon, for his valuable suggestions "
        "and constant encouragement for our project work.\n\n"
        "I also sincerely acknowledge the help and cooperation from the teaching and non-teaching staff of the Department of Computer Engineering, Government Residence Women’s "
        "Polytechnic, Tasgaon, and my mentor Dr. Vinod R. Falmari for their continuous guidance, technical support, and motivation throughout this internship period.\n\n"
        "Finally, I thank my family and friends for their unwavering support and encouragement during the completion of this project."
    )
    p = doc.add_paragraph(ack_text)
    set_run_style(p.runs[0], font_size=12)
    
    p = doc.add_paragraph("\nPlace: Tasgaon\nDate: __/__/2026\n\nAmruta Chandrakant Shinde\nRoll No: 30")
    set_run_style(p.runs[0], font_size=12)
    doc.add_page_break()

    # =========================================================================
    # INDEX
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INDEX")
    set_run_style(run, font_size=16, bold=True)
    doc.add_paragraph()

    index_data = [
        ("Chapter 1", "Organizational Structure (Industry)", "1-3"),
        ("Chapter 2", "Introduction of Industry", "4-6"),
        ("Chapter 3", "Safety Procedures Followed and Safety Gears Used", "7-8"),
        ("Chapter 4", "Problem Statement and Requirement Analysis", "9-12"),
        ("Chapter 5", "Proposed Methodology and Planning", "13-15"),
        ("Chapter 6", "Design: Flowcharts / Use Case Diagrams / Screen Designs", "16-19"),
        ("Chapter 7", "Implementation and Testing", "20-30"),
        ("Chapter 8", "Challenges Faced and Future Scope", "31-33"),
        ("Chapter 9", "Conclusion", "34"),
        ("Chapter 10", "References and Bibliography", "35-36"),
    ]
    
    table = doc.add_table(rows=len(index_data)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Chapter No."
    hdr[1].text = "Chapter Name"
    hdr[2].text = "Page No."
    
    for i, (chap, name, page) in enumerate(index_data):
        row = table.rows[i+1].cells
        row[0].text = chap
        row[1].text = name
        row[2].text = page
        
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 1: ORGANIZATIONAL STRUCTURE
    # =========================================================================
    doc.add_heading('CHAPTER 1: ORGANIZATIONAL STRUCTURE (INDUSTRY)', level=1)
    doc.add_paragraph('Industry Name: iGAP TECHNOLOGIES PVT LTD\n')
    
    org_structure = [
        ("CEO (Chief Executive Officer)", [
            "Overall head of the company.",
            "Defines company vision, strategy, and goals.",
            "Ensures all departments align with business objectives.",
            "Makes key decisions, manages stakeholders, and represents the company externally."
        ]),
        ("Chief Operating Officer (COO)", [
            "Oversees day-to-day operations.",
            "Coordinates across departments to ensure smooth execution.",
            "Focuses on operational efficiency, performance, and productivity."
        ]),
        ("Legal Team", [
            "Handles legal compliance, contracts, intellectual property, and risk management.",
            "Advises on business deals and employment law.",
            "Protects the company against lawsuits and regulatory issues."
        ]),
        ("HR Head", [
            "Manages recruitment, training, employee engagement, payroll, and performance.",
            "Builds policies to support workplace culture.",
            "Ensures compliance with labour laws."
        ]),
        ("Chief Financial Officer (CFO)", [
            "Manages finances, budgets, investments, and financial risks.",
            "Prepares reports for decision-making.",
            "Ensures profitability and financial sustainability."
        ]),
        ("Chief Information Officer (CIO)", [
            "Responsible for technology strategy and innovation.",
            "Oversees software development, IT systems, and digital transformation.",
            "Ensures the company stays competitive with new tech."
        ])
    ]
    
    for title, points in org_structure:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_style(run, font_size=12, bold=True)
        for point in points:
            doc.add_paragraph(point, style='List Bullet')
            
    doc.add_paragraph("\nSub-departments operate under these heads, including Analysis & Design, Software Architecture, Programming, Systems Support, and Marketing, ensuring a cohesive workflow from requirement gathering to final deployment.")
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 2: INTRODUCTION OF INDUSTRY
    # =========================================================================
    doc.add_heading('CHAPTER 2: INTRODUCTION OF INDUSTRY', level=1)
    
    intro_text = (
        "iGAP Technologies Pvt Ltd is a dynamic software development and training company, providing cutting-edge solutions and comprehensive training programs. "
        "Established initially as a proprietary company in 2017, iGAP Technologies evolved into a Private Limited Company in September 2021. The company is spearheaded "
        "by two directors, Mr. Abhijit Prakash Gatade (Computer Engineer) and Mr. Dilip Prakash Gatade (MBA), who bring a wealth of knowledge in both IT and business management.\n\n"
        "Departments & Services:\n"
        "1. Web Application Development: Crafting robust, scalable, and secure web applications.\n"
        "2. Website Development: Creating engaging and user-friendly websites.\n"
        "3. Mobile Application Development: Developing intuitive applications using platforms like Flutter.\n"
        "4. Data-related Services: Providing data analysis, management, and engineering services.\n\n"
        "Training Programs:\n"
        "1. Full Stack Web Development\n2. Flutter Mobile Application Development\n3. Machine Learning Engineering\n\n"
        "Work Culture:\n"
        "- Collaborative & Supportive: Focus on teamwork and skill development.\n"
        "- Work-Life Balance: Promotes overall well-being among employees.\n"
        "- Professional Environment: Smart-casual dress code and modern workspace infrastructure."
    )
    doc.add_paragraph(intro_text)
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 3: SAFETY PROCEDURES
    # =========================================================================
    doc.add_heading('CHAPTER 3: SAFETY PROCEDURES FOLLOWED AND SAFETY GEARS USED', level=1)
    
    safety_text = (
        "To maintain a safe, secure, and professional working environment, iGAP Industry follows strict organizational rules and behavioural guidelines. "
        "These policies ensure employee safety, smooth operations, and a positive workplace culture.\n\n"
        "Safety Procedures Followed:\n"
        "1. Always follow the institution's and organization's social media usage policies.\n"
        "2. Never disclose the names or personal details of students or clients in public forums.\n"
        "3. Be cautious when sharing project or user-related information with colleagues or outsiders.\n"
        "4. Use only official and verified email addresses for communication.\n"
        "5. Understand and follow data privacy guidelines (e.g., COPPA) if working with sensitive data.\n"
        "6. Choose tools, libraries, and technologies from secure, reputable sources only.\n"
        "7. Do not share or expose sensitive data such as passwords, user information, or datasets.\n"
        "8. Keep all passwords secure—do not share them with anyone.\n"
        "9. Practice ethical digital behavior—stay respectful, responsible, and professional online.\n"
        "10. Maintain academic honesty—do not plagiarize or fake data in reports or projects.\n"
        "11. Take regular breaks during computer use to avoid eye strain and mental fatigue.\n\n"
        "Safety Gears & Ergonomics Used:\n"
        "Although a software company does not require heavy industrial protective equipment, iGAP ensures:\n"
        "- Ergonomic Equipment: Adjustable chairs, wrist supports, and anti-glare screens to prevent strain.\n"
        "- Health & Hygiene: Regular sanitization of workspaces and first-aid kits available.\n"
        "- Cybersecurity Gear: Multi-factor authentication and restricted access to sensitive data."
    )
    doc.add_paragraph(safety_text)
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 4: PROBLEM STATEMENT & REQUIREMENT ANALYSIS
    # =========================================================================
    doc.add_heading('CHAPTER 4: PROBLEM STATEMENT AND REQUIREMENT ANALYSIS', level=1)
    
    doc.add_heading('4.1 Problem Statement', level=2)
    prob_text = (
        "Students have to study from a large amount of educational material such as college notes, textbooks, lecture notes, and PDF documents. "
        "Finding a particular topic or answer from lengthy study material manually is time-consuming and difficult, especially during examination preparation.\n\n"
        "Traditional PDF readers mainly allow users to read and search documents using keywords. They do not provide contextual answers, automatic summaries, "
        "practice MCQs, or important examination questions from the student's own study material. Students therefore need to use different tools for different study activities.\n\n"
        "There is a need for an intelligent system that can understand uploaded college study material, retrieve relevant information, and provide meaningful answers "
        "to students in natural language. The AI College Assistant solves this problem by providing an AI-powered platform where students can upload college PDFs "
        "and interact with their study material using natural language. The system extracts text, divides the content into smaller chunks, generates semantic embeddings, "
        "stores them in a FAISS vector database, and uses Retrieval-Augmented Generation (RAG) to retrieve relevant information. The retrieved information is provided "
        "to the Gemini AI model to generate answers based on the uploaded document. The system also provides study summaries, MCQ generation, important examination "
        "questions, chat history, and an AI-driven study planner with progress tracking."
    )
    doc.add_paragraph(prob_text)

    doc.add_heading('4.2 Requirement Analysis', level=2)
    doc.add_heading('4.2.1 Functional Requirements', level=3)
    func_req = [
        "PDF Upload: Allow students to upload college study material in PDF format.",
        "PDF Text Extraction: Extract text from the uploaded PDF for further processing.",
        "Text Chunking: Divide the extracted document into smaller overlapping chunks to improve information retrieval.",
        "Text Embedding: Convert document chunks into numerical vector representations using a HuggingFace sentence-transformer model.",
        "Vector Storage: Store generated embeddings in a FAISS vector database.",
        "Similarity Search: Retrieve relevant document chunks according to the user's question.",
        "AI Question Answering: Generate answers using retrieved document information and Gemini AI.",
        "Study Summary: Generate concise study notes from uploaded study material.",
        "MCQ Generation: Generate multiple-choice questions based on the study material.",
        "Important Questions: Generate important questions useful for examination preparation.",
        "Chat History: Store previous questions and AI-generated answers using SQLite.",
        "Study Planner: Allow users to plan study hours according to their time with the help of AI.",
        "Progress Bar: Display task completion progress for the day."
    ]
    for req in func_req:
        doc.add_paragraph(req, style='List Number')

    doc.add_heading('4.2.2 Non-Functional Requirements', level=3)
    non_func_req = [
        "Accuracy: The system should provide answers strictly relevant to the uploaded study material.",
        "Performance: Document retrieval and AI response generation should be completed within a reasonable time.",
        "Reliability: The application should handle invalid documents, missing API keys, and API errors properly.",
        "Usability: The interface should be simple and easy for students to operate (Streamlit UI).",
        "Scalability: Additional educational features and AI models should be capable of being added in the future.",
        "Security: API keys must be stored securely and should not be exposed in the source code."
    ]
    for req in non_func_req:
        doc.add_paragraph(req, style='List Number')

    doc.add_heading('4.2.3 Hardware & Software Requirements', level=3)
    hw_sw = (
        "Hardware Requirements:\n"
        "- Laptop/Desktop Computer\n- Minimum 8 GB RAM\n- Minimum 10 GB available storage\n- Internet connection for Gemini API\n\n"
        "Software Requirements:\n"
        "- Programming Language: Python\n"
        "- Framework: Streamlit\n"
        "- Libraries: PyMuPDF, LangChain, HuggingFace Sentence Transformers, FAISS, Google GenAI, SQLite, Streamlit Option Menu\n"
        "- Operating System: Windows\n"
        "- Development Environment: VS Code / Python IDE"
    )
    doc.add_paragraph(hw_sw)
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 5: PROPOSED METHODOLOGY AND PLANNING
    # =========================================================================
    doc.add_heading('CHAPTER 5: PROPOSED METHODOLOGY AND PLANNING', level=1)
    doc.add_paragraph("The project was planned and developed in different phases to ensure systematic development of the AI College Assistant over 12 weeks.")
    
    planning_data = [
        ("Week 1", "Research and Project Selection: Study problems faced by students, research existing AI educational apps, check feasibility."),
        ("Week 2", "Project Finalisation: Finalize objectives, scope, features, and expected outcomes."),
        ("Week 3", "Project Planning: Divide project into modules (PDF processing, RAG, AI features, vector DB, chat history)."),
        ("Week 4", "Synopsis Preparation: Prepare introduction, problem statement, technologies, and present synopsis."),
        ("Week 5", "Requirement Gathering: Identify functional/non-functional requirements, finalize hardware/software."),
        ("Week 6", "Environment Setup: Install Python, Streamlit, required libraries, configure Gemini API."),
        ("Week 7", "System Design: Design system architecture, workflow, and database schema."),
        ("Week 8", "Implementation: Implement PDF extraction, chunking, embeddings, FAISS, RAG, Gemini integration."),
        ("Week 9", "Testing: Test PDF processing, AI Chat, Summary, MCQ Generator, Important Questions, Chat History."),
        ("Week 10", "Documentation: Prepare system documentation, screenshots, testing results, challenges, future scope."),
        ("Week 11", "Presentation: Prepare PPT and demonstrate project features, architecture, technologies, and results."),
        ("Week 12", "Final Submission: Submit final report, source code, presentation, and deploy working application.")
    ]
    
    table = doc.add_table(rows=len(planning_data)+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Timeframe"
    hdr[1].text = "Activities"
    
    for i, (week, activity) in enumerate(planning_data):
        row = table.rows[i+1].cells
        row[0].text = week
        row[1].text = activity
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 6: DESIGN
    # =========================================================================
    doc.add_heading('CHAPTER 6: DESIGN: FLOWCHARTS / USE CASE / SCREEN DESIGNS', level=1)
    
    doc.add_heading('6.1 System Architecture & Flowchart Description', level=2)
    flow_desc = (
        "1. Start: The application is initialized using Streamlit.\n"
        "2. Upload PDF: The user uploads college study material through the sidebar.\n"
        "3. Extract Text: The system extracts readable text from the uploaded PDF using PyMuPDF.\n"
        "4. Text Chunking: The extracted text is divided into smaller overlapping chunks (e.g., 1000 chars, 200 overlap).\n"
        "5. Embedding Generation: Each text chunk is converted into a vector representation using HuggingFace sentence-transformers.\n"
        "6. FAISS Vector Store: The generated embeddings are stored in FAISS for fast similarity search.\n"
        "7. Ask Question: The user enters a question related to the uploaded study material.\n"
        "8. Similarity Search: FAISS searches for the most relevant document chunks.\n"
        "9. Context Retrieval: The retrieved chunks are combined as context.\n"
        "10. Gemini AI (RAG): The context and user question are provided to Gemini through the RAG process.\n"
        "11. Generate Answer: Gemini generates a natural-language answer based on the retrieved document information.\n"
        "12. Display & Save: The answer is displayed, and the Q&A pair is saved to the SQLite database."
    )
    doc.add_paragraph(flow_desc)
    
    add_image_placeholder(doc, "Figure 1: System Architecture Diagram", 3.0)
    add_image_placeholder(doc, "Figure 2: Stepwise Flowchart", 3.0)

    doc.add_heading('6.2 Use Case Example', level=2)
    use_case = (
        "Scenario: A student wants to understand a topic from an uploaded college PDF.\n"
        "Actor: College Student\n"
        "System: AI College Assistant\n"
        "Precondition: A valid college PDF has been uploaded and processed.\n"
        "Main Success Scenario:\n"
        "1. The student uploads a college PDF.\n"
        "2. The system extracts the text and divides it into chunks.\n"
        "3. Embeddings are generated and stored in FAISS.\n"
        "4. The student opens the AI Chat section and enters a question.\n"
        "5. FAISS retrieves relevant information from the document.\n"
        "6. The retrieved context is sent to Gemini.\n"
        "7. Gemini generates an answer, which is displayed to the student and saved in chat history."
    )
    doc.add_paragraph(use_case)

    doc.add_heading('6.3 Screen Designs', level=2)
    add_image_placeholder(doc, "Figure 3: Main Dashboard / Sidebar with PDF Upload", 3.5)
    add_image_placeholder(doc, "Figure 4: AI Chat Interface showing Q&A", 3.5)
    add_image_placeholder(doc, "Figure 5: Study Summary and MCQ Generator Output", 3.5)
    add_image_placeholder(doc, "Figure 6: Study Planner and Progress Bar Interface", 3.5)
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 7: IMPLEMENTATION AND TESTING
    # =========================================================================
    doc.add_heading('CHAPTER 7: IMPLEMENTATION AND TESTING', level=1)
    
    doc.add_heading('7.1 Implementation Details', level=2)
    impl_text = (
        "The AI College Assistant was developed using Python and the Streamlit framework. Key implementation steps include:\n\n"
        "1. PDF Processing: PyMuPDF extracts text, which is stored in the Streamlit session state.\n"
        "2. Text Chunking: Recursive text splitter divides documents (Chunk size: 1000, Overlap: 200) to preserve context.\n"
        "3. Embedding Generation: HuggingFace 'all-MiniLM-L6-v2' generates embeddings for semantic search.\n"
        "4. Vector Store: FAISS stores embeddings and performs fast similarity searches.\n"
        "5. RAG Pipeline: Combines retrieval and generation. Prompt instructs Gemini to use ONLY the retrieved college notes.\n"
        "6. Study Aids: Separate prompt templates are used for generating Summaries, MCQs, and Important Questions.\n"
        "7. Database: SQLite stores chat history (Question, Answer, Timestamp) for persistence."
    )
    doc.add_paragraph(impl_text)

    doc.add_heading('7.2 Code Snippets', level=2)
    code_text = (
        "# Example: Text Chunking and Embedding\n"
        "from langchain.text_splitter import RecursiveCharacterTextSplitter\n"
        "from langchain_huggingface import HuggingFaceEmbeddings\n"
        "import faiss\n\n"
        "text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)\n"
        "chunks = text_splitter.split_text(extracted_text)\n\n"
        "embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')\n"
        "# FAISS index creation and storage logic follows..."
    )
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    doc.add_heading('7.3 Testing Strategy & Test Cases', level=2)
    test_cases = [
        ("PDF Upload", "Valid PDF file", "PDF uploaded and processed successfully", "Pass"),
        ("Text Extraction", "Text-based PDF", "Text extracted without errors", "Pass"),
        ("Empty PDF", "PDF without readable text", "Appropriate error message displayed", "Pass"),
        ("Chunking", "Extracted PDF text", "Text divided into overlapping chunks", "Pass"),
        ("FAISS Creation", "Text chunks", "Vector store created successfully", "Pass"),
        ("AI Chat", "Question from PDF", "Relevant, context-aware AI answer generated", "Pass"),
        ("Invalid Question", "Empty or irrelevant query", "Warning displayed asking for valid input", "Pass"),
        ("Summary", "Uploaded study material", "Concise summary generated", "Pass"),
        ("MCQ Generation", "Uploaded study material", "Practice MCQs generated correctly", "Pass"),
        ("Important Questions", "Uploaded study material", "Exam-focused questions generated", "Pass"),
        ("Chat History", "Question and answer pair", "Conversation saved to SQLite", "Pass"),
        ("Clear History", "Existing history", "History deleted successfully", "Pass"),
        ("Study Planner", "Subject and hours input", "Personalized schedule generated", "Pass")
    ]
    
    table = doc.add_table(rows=len(test_cases)+1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Test Case"
    hdr[1].text = "Input"
    hdr[2].text = "Expected Result"
    hdr[3].text = "Status"
    
    for i, (tc, inp, exp, status) in enumerate(test_cases):
        row = table.rows[i+1].cells
        row[0].text = tc
        row[1].text = inp
        row[2].text = exp
        row[3].text = status
        
    doc.add_paragraph("\nNote: AI-generation features depend on Gemini API availability and quota limits.")
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 8: CHALLENGES FACED AND FUTURE SCOPE
    # =========================================================================
    doc.add_heading('CHAPTER 8: CHALLENGES FACED AND FUTURE SCOPE', level=1)
    
    doc.add_heading('8.1 Challenges Faced', level=2)
    challenges = (
        "Technical Challenges:\n"
        "- Integrating multiple AI and NLP libraries (LangChain, FAISS, HuggingFace) into a single Streamlit application.\n"
        "- Selecting the appropriate chunk size and overlap to maintain context without exceeding token limits.\n"
        "- Managing Streamlit session state to prevent re-uploading and re-processing of PDFs on every interaction.\n"
        "- Handling API errors, rate limits, and quota limitations from the Gemini API.\n\n"
        "Data-Related Challenges:\n"
        "- Extracting clean text from complex PDF formats (e.g., multi-column layouts, tables).\n"
        "- Handling PDFs containing images or scanned pages (requires future OCR integration).\n"
        "- Ensuring that retrieved chunks are highly relevant to the student's specific question.\n\n"
        "System Challenges:\n"
        "- Internet connectivity is strictly required for Gemini API requests and embedding model downloads.\n"
        "- Processing very large documents can require additional time and memory."
    )
    doc.add_paragraph(challenges)

    doc.add_heading('8.2 Future Scope', level=2)
    future_scope = (
        "1. OCR Support: Add Optical Character Recognition (OCR) to process scanned PDFs and image-based notes.\n"
        "2. Advanced Retrieval: Implement hybrid search (keyword + semantic) and document re-ranking for higher accuracy.\n"
        "3. Multi-Document Support: Allow students to upload and query multiple PDFs simultaneously, organized by subject.\n"
        "4. User Authentication: Add login functionality to maintain personalized study profiles and isolated chat histories.\n"
        "5. Multilingual Support: Enable question answering and summary generation in regional languages (e.g., Marathi, Hindi).\n"
        "6. Export Features: Allow users to download generated summaries and MCQs as PDF or Word documents.\n"
        "7. Voice Integration: Add speech-to-text for asking questions and text-to-speech for reading answers aloud.\n"
        "8. Broader Applications: Adapt the system for competitive examination preparation, institutional knowledge bases, and corporate training materials."
    )
    doc.add_paragraph(future_scope)
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 9: CONCLUSION
    # =========================================================================
    doc.add_heading('CHAPTER 9: CONCLUSION', level=1)
    conclusion = (
        "The 'AI College Assistant + Study Planner' was developed to address the difficulties students face while studying from lengthy college notes and PDF documents. "
        "The system provides an interactive, intelligent platform that allows students to upload study material and obtain useful, context-aware information through Artificial Intelligence.\n\n"
        "The project successfully combines PDF text extraction, text chunking, HuggingFace embeddings, FAISS vector search, Retrieval-Augmented Generation (RAG), Google Gemini AI, "
        "Streamlit, and SQLite to create a cohesive educational assistant. The RAG-based AI Chat feature ensures that the system provides document-based responses rather than "
        "relying solely on general AI knowledge, significantly reducing hallucinations and improving accuracy.\n\n"
        "Beyond simple Q&A, the project provides additional educational features such as Study Summary, MCQ Generator, Important Questions, Chat History, and an AI-driven Study Planner. "
        "These features help students understand, revise, and practice their study material from a single, unified application.\n\n"
        "During implementation, challenges such as document processing, vector retrieval, API integration, and deployment were systematically addressed. Testing confirmed that all major "
        "application modules work together successfully. Overall, the AI College Assistant demonstrates how Generative AI, NLP, semantic search, and vector databases can be combined "
        "to develop a highly practical, impactful educational application. The system provides a strong foundation for future development into a comprehensive, personalized AI-powered learning platform."
    )
    doc.add_paragraph(conclusion)
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 10: REFERENCES AND BIBLIOGRAPHY
    # =========================================================================
    doc.add_heading('CHAPTER 10: REFERENCES AND BIBLIOGRAPHY', level=1)
    
    refs = [
        "1. Python Software Foundation, Python Documentation. https://docs.python.org",
        "2. Streamlit, Streamlit Documentation. https://docs.streamlit.io",
        "3. LangChain, LangChain Documentation. https://python.langchain.com",
        "4. Meta AI Research, FAISS – Facebook AI Similarity Search. https://faiss.ai",
        "5. Hugging Face, Sentence Transformers Documentation. https://www.sbert.net",
        "6. Google, Gemini API Documentation. https://ai.google.dev/gemini-api/docs",
        "7. PyMuPDF, PyMuPDF Documentation. https://pymupdf.readthedocs.io",
        "8. SQLite, SQLite Documentation. https://www.sqlite.org/docs.html",
        "9. Lewis, P. et al., 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks', 2020.",
        "10. Reimers, N. and Gurevych, I., 'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks', 2019.",
        "11. iGAP Technologies Pvt. Ltd. Official Website and Internal Training Materials.",
        "12. Lecture notes and guidance from faculty during the industrial training period."
    ]
    
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)

    # Save the document
    file_name = "AI_College_Assistant_Internship_Report.docx"
    doc.save(file_name)
    print(f"Successfully generated: {file_name}")
    print("NOTE: To reach the full 40-45 page count, open the generated Word document and:")
    print("1. Replace the '[ INSERT IMAGE HERE ]' placeholders with actual large screenshots from your Streamlit app.")
    print("2. Ensure the document is formatted with 1.5 line spacing, 12pt Times New Roman font, and 1-inch margins.")
    print("3. Add page breaks where necessary to ensure chapters start on new pages.")

if __name__ == "__main__":
    main()