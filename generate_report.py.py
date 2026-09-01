from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_run_style(run, font_name='Times New Roman', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_image_placeholder(doc, text, height_inches=4.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(f"\n\n[ INSERT LARGE SCREENSHOT/DIAGRAM HERE ]\n\n{text}\n\n")
    runner.bold = True
    runner.font.color.rgb = RGBColor(128, 128, 128)
    runner.font.size = Pt(12)
    doc.add_paragraph()

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # =========================================================================
    # COVER PAGE & CERTIFICATES (Pages 1-4)
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Government Residence Women Polytechnic, Tasgaon\n2026 – 2027\n\n")
    set_run_style(run, font_size=16, bold=True)
    run = p.add_run("A Report on\n“INTERNSHIP”\n(ITR – 315004)\n\n")
    set_run_style(run, font_size=16, bold=True)
    run = p.add_run("Project Title:\nAI College Assistant + Study Planner\n\n")
    set_run_style(run, font_size=14, bold=True, color=RGBColor(0, 51, 102))
    run = p.add_run("Department of Computer Engineering\nMAHARASHTRA STATE BOARD OF TECHNICAL EDUCATION, MUMBAI\n\n\n\n")
    set_run_style(run, font_size=14, bold=True)
    run = p.add_run("Submitted by:\nAmruta Chandrakant Shinde\nRoll No: 30 | Enrolment No: 24210280163\n\n")
    set_run_style(run, font_size=12, bold=True)
    run = p.add_run("Under the Guidance of:\nDr. Vinod R. Falmari (Mentor)\nSmt. S. B. Patil (HOD, Computer Engineering)\nDr. S. B. Patil (Principal, GRWPT)")
    set_run_style(run, font_size=12, bold=True)
    doc.add_page_break()

    # Certificate & Industrial Certificate & Abstract & Acknowledgement
    for title in ["CERTIFICATE", "INDUSTRIAL CERTIFICATE", "ABSTRACT", "ACKNOWLEDGEMENT"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_run_style(run, font_size=16, bold=True)
        doc.add_paragraph()
        if title == "ABSTRACT":
            doc.add_paragraph("The 12-week industrial training at iGAP Technologies Pvt. Ltd. provided an excellent opportunity to bridge the gap between academic learning and industrial practices. The training was designed to enhance technical skills, professional competence, and practical exposure to real-world IT projects. During the program, I gained hands-on experience in areas such as software development, data science, artificial intelligence, and machine learning using Python. The training emphasized key concepts including data preprocessing, exploratory data analysis, implementation of machine learning models, and deployment of solutions for real-time applications.\n\nIndustrial Training provides required professional and practical skills to the students. It is an essential part in the development of the practical and professional skills required of an engineer and an aid to prospective employment. For this training, I joined iGAP Technologies Pvt. Ltd., where I learned Data Science and AI/ML using Python and developed a comprehensive project.\n\nThe project developed is the 'AI College Assistant + Study Planner'. Students often struggle with lengthy educational materials like college notes, textbooks, and PDF documents. Finding specific topics or answers manually is time-consuming. Traditional PDF readers only allow keyword searches and lack contextual understanding, automatic summarization, or practice question generation. The AI College Assistant solves this by providing an intelligent platform where students can upload college PDFs and interact with the material using natural language. The system extracts text, divides it into semantic chunks, generates embeddings using HuggingFace models, and stores them in a FAISS vector database. Using Retrieval-Augmented Generation (RAG) powered by the Google Gemini AI model, the system provides accurate, context-aware answers, generates concise study summaries, creates practice MCQs, identifies important examination questions, and offers a personalized AI-driven study planner with progress tracking. This project successfully demonstrates the practical application of Generative AI and NLP in solving real-world educational challenges.")
        elif title == "ACKNOWLEDGEMENT":
            doc.add_paragraph("I would like to express my sincere gratitude to iGAP Technologies Pvt. Ltd. for providing me with the opportunity to undergo 12 weeks of industrial training in Data Science and Artificial Intelligence/Machine Learning using Python. This training has been an invaluable learning experience, allowing me to apply theoretical knowledge to practical, real-world problems and gain industry-relevant skills.\n\nIt is a great privilege for me to express our sincere thanks to Dr. S. B. Patil, Principal of Government Residence Women Polytechnic, Tasgaon, for his valuable suggestions and constant encouragement for our project work.\n\nWe sincerely acknowledge the help and cooperation from the teaching and non-teaching staff of the Department of Computer Engineering, Government Residence Women’s Polytechnic, Tasgaon, and my mentor Dr. Vinod R. Falmari for their continuous guidance, technical support, and motivation throughout this internship period.\n\nFinally, I thank my family and friends for their unwavering support and encouragement during the completion of this project.\n\nPlace: Tasgaon\nDate: __/__/2026\nAmruta Chandrakant Shinde\nRoll No: 30")
        else:
            doc.add_paragraph("This is to certify that the student mentioned below of Diploma in Computer Engineering of Government Residence Women Polytechnic, Tasgaon (1228), has satisfactorily completed a 12-week duration internship and has submitted this report as partial fulfilment of the prescribed curriculum of the Maharashtra State Board of Technical Education, Mumbai, for the academic year 2026–2027.\n\nRoll No: 30 | Enrolment No: 24210280163 | Name: Amruta Chandrakant Shinde\n\nDr. Vinod R. Falmari (Mentor) | Smt. S. B. Patil (HOD) | Dr. S. B. Patil (Principal)")
        doc.add_page_break()

    # =========================================================================
    # INDEX (Page 5)
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INDEX")
    set_run_style(run, font_size=16, bold=True)
    doc.add_paragraph()
    index_data = [("1", "Organizational Structure (Industry)", "6-9"), ("2", "Introduction of Industry", "10-12"), ("3", "Safety Procedures Followed and Safety Gears Used", "13-14"), ("4", "Problem Statement and Requirement Analysis", "15-18"), ("5", "Proposed Methodology and Planning", "19-21"), ("6", "Design: Flowcharts / Use Case / Screen Designs", "22-26"), ("7", "Implementation and Testing", "27-38"), ("8", "Challenges Faced and Future Scope", "39-41"), ("9", "Conclusion", "42"), ("10", "References and Bibliography", "43-45")]
    table = doc.add_table(rows=len(index_data)+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Chapter No.", "Chapter Name", "Page No."
    for i, (chap, name, page) in enumerate(index_data):
        row = table.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = chap, name, page
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 1 & 2 (Pages 6-12)
    # =========================================================================
    doc.add_heading('CHAPTER 1: ORGANIZATIONAL STRUCTURE (INDUSTRY)', level=1)
    doc.add_paragraph('Industry Name: iGAP TECHNOLOGIES PVT LTD\n')
    org_structure = [
        ("CEO (Chief Executive Officer)", ["Overall head of the company.", "Defines company vision, strategy, and goals.", "Ensures all departments align with business objectives.", "Makes key decisions, manages stakeholders, and represents the company externally."]),
        ("Chief Operating Officer (COO)", ["Oversees day-to-day operations.", "Coordinates across departments to ensure smooth execution.", "Focuses on operational efficiency, performance, and productivity."]),
        ("Legal Team", ["Handles legal compliance, contracts, intellectual property, and risk management.", "Advises on business deals and employment law."]),
        ("HR Head", ["Manages recruitment, training, employee engagement, payroll, and performance.", "Builds policies to support workplace culture."]),
        ("Chief Financial Officer (CFO)", ["Manages finances, budgets, investments, and financial risks.", "Prepares reports for decision-making."]),
        ("Operations Leader", ["Ensures IT infrastructure, systems, and processes are reliable.", "Manages IT support, networking, and installations."]),
        ("Chief Information Officer (CIO)", ["Responsible for technology strategy and innovation.", "Oversees software development, IT systems, and digital transformation."])
    ]
    for title, points in org_structure:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_style(run, font_size=12, bold=True)
        for point in points: doc.add_paragraph(point, style='List Bullet')
    doc.add_page_break()

    doc.add_heading('CHAPTER 2: INTRODUCTION OF INDUSTRY', level=1)
    doc.add_paragraph("iGAP Technologies Pvt Ltd is a dynamic software development and training company, providing cutting-edge solutions and comprehensive training programs. Established initially as a proprietary company in 2017, iGAP Technologies evolved into a Private Limited Company in September 2021. The company is spearheaded by two directors, Mr. Abhijit Prakash Gatade (Computer Engineer) and Mr. Dilip Prakash Gatade (MBA).\n\nDepartments & Services:\n1. Web Application Development\n2. Website Development\n3. Mobile Application Development\n4. Data-related Services\n\nTraining Programs:\n1. Full Stack Web Development\n2. Flutter Mobile Application Development\n3. Machine Learning Engineering\n\nWork Culture:\n- Collaborative & Supportive\n- Work-Life Balance\n- Professional Environment")
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 3 & 4 (Pages 13-18)
    # =========================================================================
    doc.add_heading('CHAPTER 3: SAFETY PROCEDURES FOLLOWED AND SAFETY GEARS USED', level=1)
    doc.add_paragraph("To maintain a safe, secure, and professional working environment, iGAP Industry follows strict organizational rules and behavioural guidelines.\n\nSafety Procedures Followed:\n1. Always follow the institution's and organization's social media usage policies.\n2. Never disclose the names or personal details of students or clients in public forums.\n3. Use only official and verified email addresses for communication.\n4. Choose tools, libraries, and technologies from secure, reputable sources only.\n5. Do not share or expose sensitive data such as passwords, user information, or datasets.\n6. Keep all passwords secure—do not share them with anyone.\n7. Practice ethical digital behavior—stay respectful, responsible, and professional online.\n8. Maintain academic honesty—do not plagiarize or fake data in reports or projects.\n9. Take regular breaks during computer use to avoid eye strain and mental fatigue.\n\nSafety Gears & Ergonomics Used:\n- Ergonomic Equipment: Adjustable chairs, wrist supports, and anti-glare screens.\n- Health & Hygiene: Regular sanitization of workspaces and first-aid kits available.\n- Cybersecurity Gear: Multi-factor authentication and restricted access to sensitive data.")
    doc.add_page_break()

    doc.add_heading('CHAPTER 4: PROBLEM STATEMENT AND REQUIREMENT ANALYSIS', level=1)
    doc.add_heading('4.1 Problem Statement', level=2)
    doc.add_paragraph("Students have to study from a large amount of educational material such as college notes, textbooks, lecture notes, and PDF documents. Finding a particular topic or answer from lengthy study material manually is time-consuming and difficult, especially during examination preparation. Traditional PDF readers mainly allow users to read and search documents using keywords. They do not provide contextual answers, automatic summaries, practice MCQs, or important examination questions from the student's own study material. There is therefore a need for an intelligent system that can understand uploaded college study material, retrieve relevant information, and provide meaningful answers to students in natural language.")
    
    doc.add_heading('4.2 Requirement Analysis', level=2)
    doc.add_heading('Functional Requirements', level=3)
    func_req = ["PDF Upload", "PDF Text Extraction", "Text Chunking", "Text Embedding", "Vector Storage", "Similarity Search", "AI Question Answering", "Study Summary", "MCQ Generation", "Important Questions", "Chat History", "History Management", "Error Handling", "User Interface", "Study Planner", "Progress Bar"]
    for req in func_req: doc.add_paragraph(f"{req} – The system must handle this functionality seamlessly.", style='List Number')
    
    doc.add_heading('Non-Functional Requirements', level=3)
    non_func_req = ["Accuracy", "Performance", "Reliability", "Usability", "Scalability", "Maintainability", "Security", "Portability"]
    for req in non_func_req: doc.add_paragraph(f"{req} – Ensures high quality and robust system operation.", style='List Number')
    
    doc.add_heading('Hardware & Software Requirements', level=3)
    doc.add_paragraph("Hardware: Laptop/Desktop, Min 8GB RAM, 10GB Storage, Internet.\nSoftware: Python, Streamlit, PyMuPDF, LangChain, HuggingFace, FAISS, Google GenAI, SQLite, Windows OS, VS Code.")
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 5 & 6 (Pages 19-26)
    # =========================================================================
    doc.add_heading('CHAPTER 5: PROPOSED METHODOLOGY AND PLANNING', level=1)
    doc.add_paragraph("The project was planned and developed in different phases to ensure systematic development of the AI College Assistant over 12 weeks.")
    planning_data = [("Week 1", "Research and Project Selection"), ("Week 2", "Project Finalisation"), ("Week 3", "Project Planning"), ("Week 4", "Synopsis Preparation"), ("Week 5", "Requirement Gathering"), ("Week 6", "Environment Setup"), ("Week 7", "System Design"), ("Week 8", "Implementation"), ("Week 9", "Testing"), ("Week 10", "Documentation"), ("Week 11", "Presentation"), ("Week 12", "Final Submission")]
    table = doc.add_table(rows=len(planning_data)+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Timeframe", "Activities"
    for i, (week, activity) in enumerate(planning_data):
        row = table.rows[i+1].cells
        row[0].text, row[1].text = week, activity
    doc.add_page_break()

    doc.add_heading('CHAPTER 6: DESIGN: FLOWCHARTS / USE CASE / SCREEN DESIGNS', level=1)
    doc.add_heading('6.1 System Architecture & Flowchart Description', level=2)
    doc.add_paragraph("1. Start: The application is initialized using Streamlit.\n2. Upload PDF: The user uploads college study material.\n3. Extract Text: PyMuPDF extracts readable text.\n4. Text Chunking: Text is divided into smaller overlapping chunks.\n5. Embedding Generation: HuggingFace converts chunks to vectors.\n6. FAISS Vector Store: Embeddings are stored in FAISS.\n7. Ask Question: User enters a question.\n8. Similarity Search: FAISS searches for relevant chunks.\n9. Context Retrieval: Chunks are combined as context.\n10. Gemini AI (RAG): Context and question are provided to Gemini.\n11. Generate Answer: Gemini generates a natural-language answer.\n12. Display & Save: Answer is displayed and saved to SQLite.")
    
    add_image_placeholder(doc, "Fig 1: System Architectural Diagram")
    doc.add_page_break()
    add_image_placeholder(doc, "Fig 2: Flowchart: Stepwise Working of the System")
    doc.add_page_break()
    add_image_placeholder(doc, "Fig 3: Use Case Diagram")
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 7: IMPLEMENTATION AND TESTING (Pages 27-38) - MASSIVE EXPANSION
    # =========================================================================
    doc.add_heading('CHAPTER 7: IMPLEMENTATION AND TESTING', level=1)
    doc.add_heading('7.1 Implementation Details', level=2)
    doc.add_paragraph("The AI College Assistant was implemented using Python and Streamlit. Key modules include:\n1. PDF Processing (PyMuPDF)\n2. Text Chunking (LangChain RecursiveTextSplitter)\n3. Embedding Generation (HuggingFace all-MiniLM-L6-v2)\n4. Vector Store (FAISS)\n5. RAG Pipeline & AI Generation (Google Gemini)\n6. Database (SQLite for Chat History)\n7. Study Planner & Progress Tracking")
    
    doc.add_heading('7.2 Information of Libraries used in Project', level=2)
    libs = [
        ("1. streamlit", "Used to build the interactive web dashboard (UI). Provides input forms, file upload, and display elements. Makes the app run as a local web app."),
        ("2. PyMuPDF (fitz)", "Used for extracting text from uploaded PDF documents efficiently and accurately."),
        ("3. LangChain", "Provides the framework for building the RAG pipeline, including text splitting and document loaders."),
        ("4. HuggingFace Sentence Transformers", "Used to convert text chunks into numerical vector representations (embeddings) for semantic search."),
        ("5. FAISS (Facebook AI Similarity Search)", "A highly efficient vector database used to store embeddings and perform fast similarity searches."),
        ("6. Google GenAI", "Integrates the Gemini AI model to generate natural language answers, summaries, and MCQs based on retrieved context."),
        ("7. SQLite3", "A lightweight relational database used to store and manage user chat history locally."),
        ("8. streamlit-option-menu", "Used to create a clean, interactive sidebar navigation menu for the application.")
    ]
    for lib_name, lib_desc in libs:
        p = doc.add_paragraph()
        run = p.add_run(lib_name)
        set_run_style(run, bold=True)
        p.add_run(f"\n{lib_desc}")
    doc.add_page_break()

    doc.add_heading('7.3 How It Works (Step-by-Step)', level=2)
    how_it_works = [
        "1. User Registration / PDF Upload: The user uploads a college PDF via the sidebar.",
        "2. Input Processing: The system extracts text and divides it into semantic chunks.",
        "3. AI-Based Vector Storage: Embeddings are generated and stored in FAISS.",
        "4. Query Processing: The user asks a question or requests a summary/MCQ.",
        "5. RAG Retrieval: FAISS retrieves the most relevant chunks.",
        "6. Result & Advice Generation: Gemini generates the final output strictly based on the PDF.",
        "7. Study Planner Integration: The system schedules tasks and tracks progress via a progress bar."
    ]
    for step in how_it_works: doc.add_paragraph(step)
    doc.add_page_break()

    doc.add_heading('7.4 Testing Section', level=2)
    doc.add_paragraph("Testing Strategy: Unit Testing, Integration Testing, System Testing, Usability Testing.")
    test_cases = [("PDF Upload", "Valid PDF", "PDF uploaded successfully", "Pass"), ("Text Extraction", "Text-based PDF", "Text extracted successfully", "Pass"), ("Empty PDF", "PDF without text", "Error message displayed", "Pass"), ("AI Chat", "Question from PDF", "Relevant AI answer generated", "Pass"), ("Summary", "Uploaded material", "Summary generated", "Pass"), ("MCQ", "Uploaded material", "MCQs generated", "Pass")]
    table = doc.add_table(rows=len(test_cases)+1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Test Case", "Input", "Expected Output", "Status"
    for i, (tc, inp, exp, status) in enumerate(test_cases):
        row = table.rows[i+1].cells
        row[0].text, row[1].text, row[2].text, row[3].text = tc, inp, exp, status
    doc.add_page_break()

    doc.add_heading('7.5 Project Outputs (Screenshots)', level=2)
    outputs = ["Fig 4: Main Dashboard & Sidebar Navigation", "Fig 5: PDF Upload and Processing Screen", "Fig 6: AI Chat Interface - Asking a Question", "Fig 7: AI Chat Interface - Generated Answer", "Fig 8: Study Summary Generation Output", "Fig 9: MCQ Generator Output", "Fig 10: Important Questions Output", "Fig 11: Chat History Screen", "Fig 12: Study Planner Interface", "Fig 13: Progress Bar and Task Management"]
    for out in outputs:
        add_image_placeholder(doc, out)
        doc.add_page_break()

    doc.add_heading('7.6 Advantages, Limitations & Skills Developed', level=2)
    doc.add_paragraph("Advantages:\n1. Reduces time required to search and understand college study material.\n2. Provides multiple study-support features through a single platform.\n3. Context-aware answers reduce AI hallucinations.\n4. User-friendly interface suitable for all students.\n\nLimitations:\n1. Gemini API requests are subject to API quota limits.\n2. Internet connectivity is required for AI features.\n3. Scanned/image-only PDFs may not provide extractable text without OCR.\n\nSkills Developed:\n1. Proper Time Management and Task Management.\n2. Quick Learning Ability and Problem Solving.\n3. Practical exposure to RAG, Vector Databases, and Generative AI.\n4. Confidence around AI development and cloud deployment.")
    doc.add_page_break()

    # =========================================================================
    # CHAPTER 8, 9, 10 (Pages 39-45)
    # =========================================================================
    doc.add_heading('CHAPTER 8: CHALLENGES FACED AND FUTURE SCOPE', level=1)
    doc.add_heading('Challenges Faced', level=2)
    doc.add_paragraph("Technical: Integrating multiple AI/NLP libraries, selecting chunk size/overlap, managing Streamlit session state, handling API quota limits.\nData-Related: Extracting clean text from complex PDFs, maintaining context, handling scanned pages.\nSystem: API quota limitations, internet dependency, processing large PDFs.")
    
    doc.add_heading('Future Scope', level=2)
    doc.add_paragraph("1. Add OCR support for scanned PDF documents.\n2. Support multiple PDFs simultaneously organized by subject.\n3. Add user login and personalized study profiles.\n4. Add downloadable PDF/Word summaries and interactive MCQ tests.\n5. Add voice-based and multilingual question answering.\n6. Expand to competitive examination preparation and institutional knowledge bases.")
    doc.add_page_break()

    doc.add_heading('CHAPTER 9: CONCLUSION', level=1)
    doc.add_paragraph("The AI College Assistant was developed to address the difficulties students face while studying from lengthy college notes and PDF documents. The system provides an interactive platform that allows students to upload study material and obtain useful information through Artificial Intelligence. The project combines PDF text extraction, text chunking, HuggingFace embeddings, FAISS vector search, Retrieval-Augmented Generation (RAG), Gemini AI, Streamlit, and SQLite to create an integrated educational assistant. The RAG-based AI Chat feature retrieves relevant information from the uploaded study material before generating an answer, ensuring document-based responses. Overall, the AI College Assistant demonstrates how Generative AI, NLP, semantic search, vector databases, and RAG can be combined to develop a practical educational application.")
    doc.add_page_break()

    doc.add_heading('CHAPTER 10: REFERENCES AND BIBLIOGRAPHY', level=1)
    refs = [
        "1. Python Software Foundation, Python Documentation. https://docs.python.org",
        "2. Streamlit, Streamlit Documentation. https://docs.streamlit.io",
        "3. LangChain, LangChain Documentation. https://python.langchain.com",
        "4. Meta AI Research, FAISS – Facebook AI Similarity Search. https://faiss.ai",
        "5. Hugging Face, Sentence Transformers Documentation. https://www.sbert.net",
        "6. Google, Gemini API Documentation. https://ai.google.dev/gemini-api/docs",
        "7. PyMuPDF, PyMuPDF Documentation. https://pymupdf.readthedocs.io",
        "8. SQLite, SQLite Documentation. https://www.sqlite.org/docs.html",
        "9. Lewis, P. et al., Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks, 2020.",
        "10. Reimers, N. and Gurevych, I., Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks, 2019.",
        "11. iGAP Technologies Pvt. Ltd. Official Website and Internal Training Materials.",
        "12. Lecture notes and guidance from faculty during the industrial training period."
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)

    file_name = "AI_College_Assistant_Internship_Report_45Pages.docx"
    doc.save(file_name)
    print(f"Successfully generated: {file_name}")

if __name__ == "__main__":
    main()